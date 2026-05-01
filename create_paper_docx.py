import re
from docx import Document

def create_doc():
    with open('DL_Academic_Paper.md', 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    sections = md_text.split('\n')
    
    for line in sections:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('*   **') or line.startswith('-   **') or line.startswith('* **'):
            p = doc.add_paragraph(style='List Bullet')
            clean_line = line.replace('*', '').strip()
            p.add_run(clean_line)
        elif line.startswith('* ') or line.startswith('- '):
             p = doc.add_paragraph(style='List Bullet')
             p.add_run(line[2:].strip())
        else:
            p = doc.add_paragraph()
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(line)
                
    doc.save('DL_Academic_Paper.docx')

if __name__ == '__main__':
    create_doc()
