import re
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    import sys
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_doc():
    with open('DL_Project_Report.md', 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    sections = md_text.split('\n')
    
    for line in sections:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('* **') or line.startswith('- **'):
            p = doc.add_paragraph(style='List Bullet')
            # very basic bold parser for lists
            clean_line = line[2:].strip()
            if '**' in clean_line:
                parts = clean_line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1: # Odd elements were inside **
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(clean_line)
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
            p = doc.add_paragraph(style='List Number')
            clean_line = line[3:].strip()
            if '**' in clean_line:
                parts = clean_line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(clean_line)
        else:
            p = doc.add_paragraph()
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(line)
                
    doc.save('DL_Project_Report.docx')
    print("Report saved as DL_Project_Report.docx")

if __name__ == '__main__':
    create_doc()
